#!/usr/bin/env python3
import csv
import json
import re
import subprocess
import sys
import zipfile
import xml.etree.ElementTree as ET
from pathlib import Path

try:
    from docx import Document
except Exception:
    Document = None

NS = {
    'a': 'http://schemas.openxmlformats.org/spreadsheetml/2006/main',
    'r': 'http://schemas.openxmlformats.org/officeDocument/2006/relationships',
}

KEYWORDS = {
    'code': ['合同编号', '编号'],
    'name': ['合同名称'],
    'partnerName': ['乙方', '供应商', '客户', '对方', '合作方'],
    'ourSide': ['甲方', '采购方', '需方'],
    'amount': ['合同金额', '总金额', '含税金额', '价税合计', '金额'],
    'taxRate': ['税率'],
    'signDate': ['签订日期', '签署日期'],
    'effectiveDate': ['生效日期'],
    'expireDate': ['到期日期', '截止日期', '终止日期'],
    'projectName': ['项目名称'],
    'summary': ['合同内容', '主要内容', '备注', '摘要'],
    'templateName': ['合同模板', '模板名称'],
    'type': ['合同类型'],
    'status': ['状态'],
    'riskLevel': ['风险等级'],
    'archiveStatus': ['归档状态'],
    'paymentStatus': ['付款状态'],
    'signingMethod': ['签署方式', '签约方式'],
}

PAYMENT_HEADERS = {
    'phase': ['付款节点', '付款阶段', '期次', '节点'],
    'type': ['收付类型', '类型'],
    'amount': ['节点金额', '付款金额', '金额'],
    'planDate': ['计划付款日', '计划日期', '付款日期'],
    'percent': ['付款比例', '比例'],
}

MILESTONE_HEADERS = {
    'name': ['履约节点', '里程碑', '交付节点'],
    'owner': ['负责人', '责任人'],
    'dueDate': ['计划完成时间', '计划日期', '截止日期'],
    'status': ['履约状态', '状态'],
}


def normalize_key(value: str) -> str:
    return re.sub(r'\s+', '', str(value or '')).strip().lower()


def first_present(source, names):
    normalized = {normalize_key(k): v for k, v in source.items()}
    for name in names:
        key = normalize_key(name)
        if key in normalized and str(normalized[key]).strip():
            return str(normalized[key]).strip()
    return ''


def read_csv(path: Path):
    with path.open('r', encoding='utf-8-sig', newline='') as f:
        return list(csv.reader(f))


def read_xlsx(path: Path):
    with zipfile.ZipFile(path) as zf:
        shared = []
        if 'xl/sharedStrings.xml' in zf.namelist():
            root = ET.fromstring(zf.read('xl/sharedStrings.xml'))
            for si in root.findall('a:si', NS):
                texts = [node.text or '' for node in si.findall('.//a:t', NS)]
                shared.append(''.join(texts))

        workbook = ET.fromstring(zf.read('xl/workbook.xml'))
        sheets = workbook.find('a:sheets', NS)
        first_sheet = sheets[0]
        rel_id = first_sheet.attrib.get('{%s}id' % NS['r'])

        rels = ET.fromstring(zf.read('xl/_rels/workbook.xml.rels'))
        target = None
        for rel in rels:
            if rel.attrib.get('Id') == rel_id:
                target = rel.attrib.get('Target')
                break
        if not target:
            raise RuntimeError('Cannot resolve first worksheet')
        if not target.startswith('worksheets/'):
            target = target.replace('../', '')
        sheet = ET.fromstring(zf.read(f'xl/{target}'))
        rows = []
        for row in sheet.findall('.//a:sheetData/a:row', NS):
            values = []
            for c in row.findall('a:c', NS):
                cell_type = c.attrib.get('t')
                v = c.find('a:v', NS)
                value = '' if v is None or v.text is None else v.text
                if cell_type == 's' and value:
                    idx = int(value)
                    value = shared[idx] if idx < len(shared) else ''
                values.append(value)
            rows.append(values)
        return rows


def rows_to_records(rows):
    if not rows:
        return []
    header = rows[0]
    records = []
    for row in rows[1:]:
        if not any(str(v).strip() for v in row):
            continue
        padded = row + [''] * (len(header) - len(row))
        records.append({str(header[i]).strip(): padded[i] for i in range(len(header))})
    return records


def parse_payment_plan_from_text(text: str):
    plans = []
    patterns = [
        r'(首付款|预付款|进度款|尾款|质保金)[^\n]{0,20}?([0-9]+(?:\.[0-9]{1,2})?)\s*元[^\n]{0,20}?(20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2})?',
        r'第([一二三四五六七八九十0-9]+)期[^\n]{0,20}?([0-9]+(?:\.[0-9]{1,2})?)\s*元[^\n]{0,20}?(20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2})?',
    ]
    for pattern in patterns:
        for match in re.finditer(pattern, text):
            phase = match.group(1)
            amount = match.group(2)
            plan_date = (match.group(3) or '').replace('年', '-').replace('月', '-').replace('日', '').replace('/', '-').replace('.', '-')
            plans.append({
                'id': f'plan-{len(plans)+1}',
                'phase': phase,
                'type': '应收',
                'amount': amount,
                'planDate': plan_date,
                'percent': ''
            })
    return plans[:6]


def parse_milestones_from_text(text: str):
    milestones = []
    keywords = ['交付', '验收', '上线', '到货', '签署', '归档', '实施']
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        if any(keyword in line for keyword in keywords):
            date = ''
            found = re.search(r'(20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2})', line)
            if found:
                date = found.group(1).replace('年', '-').replace('月', '-').replace('日', '').replace('/', '-').replace('.', '-')
            milestones.append({
                'id': f'ms-{len(milestones)+1}',
                'name': line[:40],
                'owner': '',
                'dueDate': date,
                'status': '未开始'
            })
    return milestones[:6]


def map_excel_record(record):
    mapped = {
        'type': first_present(record, KEYWORDS['type']) or '销售',
        'status': first_present(record, KEYWORDS['status']) or '草稿',
        'currency': 'CNY',
        'riskLevel': first_present(record, KEYWORDS['riskLevel']) or '正常',
        'performanceStatus': '未开始',
        'archiveStatus': first_present(record, KEYWORDS['archiveStatus']) or '未归档',
        'paymentStatus': first_present(record, KEYWORDS['paymentStatus']) or '待维护',
        'templateName': first_present(record, KEYWORDS['templateName']),
        'signingMethod': first_present(record, KEYWORDS['signingMethod']) or '线下签署',
    }
    for field, candidates in KEYWORDS.items():
        value = first_present(record, candidates)
        if value:
            mapped[field] = value

    payment = {field: first_present(record, names) for field, names in PAYMENT_HEADERS.items()}
    if any(payment.values()):
        mapped['paymentPlan'] = [{
            'id': 'plan-1',
            'phase': payment.get('phase') or '付款节点1',
            'type': payment.get('type') or '应收',
            'amount': payment.get('amount') or '',
            'planDate': payment.get('planDate') or '',
            'percent': payment.get('percent') or '',
        }]

    milestone = {field: first_present(record, names) for field, names in MILESTONE_HEADERS.items()}
    if any(milestone.values()):
        mapped['milestones'] = [{
            'id': 'ms-1',
            'name': milestone.get('name') or '履约节点1',
            'owner': milestone.get('owner') or '',
            'dueDate': milestone.get('dueDate') or '',
            'status': milestone.get('status') or '未开始',
        }]
    return mapped


def extract_text_pairs(text: str):
    result = {}
    lines = [line.strip() for line in text.splitlines() if line.strip()]
    for line in lines:
        for field, keywords in KEYWORDS.items():
            for keyword in keywords:
                pattern = rf'{re.escape(keyword)}[：: ]+(.+)$'
                match = re.search(pattern, line)
                if match:
                    result[field] = match.group(1).strip()
                    break
            if field in result:
                break

    amount_match = re.search(r'([0-9]+(?:\.[0-9]{1,2})?)\s*元', text)
    if amount_match and 'amount' not in result:
        result['amount'] = amount_match.group(1)
    tax_match = re.search(r'税率[：: ]*([0-9]+)', text)
    if tax_match and 'taxRate' not in result:
        result['taxRate'] = tax_match.group(1)
    dates = re.findall(r'20\d{2}[-/.年]\d{1,2}[-/.月]\d{1,2}', text)
    if dates and 'signDate' not in result:
        result['signDate'] = dates[0].replace('年', '-').replace('月', '-').replace('日', '').replace('/', '-').replace('.', '-')

    result['paymentPlan'] = parse_payment_plan_from_text(text)
    result['milestones'] = parse_milestones_from_text(text)
    return result


def parse_docx(path: Path):
    if Document is None:
        raise RuntimeError('python-docx unavailable')
    doc = Document(str(path))
    texts = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            texts.append(' | '.join([c for c in cells if c]))
    full_text = '\n'.join(texts)
    extracted = extract_text_pairs(full_text)
    extracted['summary'] = '\n'.join(texts[:8])[:400]
    return extracted


def parse_text(path: Path):
    text = path.read_text(encoding='utf-8', errors='ignore')
    extracted = extract_text_pairs(text)
    extracted['summary'] = text[:400]
    return extracted


def parse_pdf(path: Path):
    proc = subprocess.run(['pdftotext', str(path), '-'], capture_output=True, text=True)
    if proc.returncode != 0:
        raise RuntimeError(proc.stderr.strip() or 'pdftotext failed')
    text = proc.stdout.strip()
    if not text:
        raise RuntimeError('PDF has no extractable text; OCR not available on this machine yet')
    extracted = extract_text_pairs(text)
    extracted['summary'] = text[:400]
    return extracted


def parse_image(path: Path):
    raise RuntimeError('Image OCR is not available yet on this machine; install tesseract or add OCR service support')


def apply_defaults(data):
    data.setdefault('type', '销售')
    data.setdefault('status', '草稿')
    data.setdefault('currency', 'CNY')
    data.setdefault('riskLevel', '正常')
    data.setdefault('performanceStatus', '未开始')
    data.setdefault('archiveStatus', '未归档')
    data.setdefault('paymentStatus', '待维护')
    data.setdefault('signingMethod', '线下签署')
    data.setdefault('paymentPlan', [])
    data.setdefault('milestones', [])
    return data


def main():
    if len(sys.argv) < 3:
        print(json.dumps({'error': 'usage: parse_import.py <excel|contract|ocr> <path>'}, ensure_ascii=False))
        sys.exit(1)

    mode = sys.argv[1]
    path = Path(sys.argv[2])
    if mode == 'excel':
        rows = read_csv(path) if path.suffix.lower() == '.csv' else read_xlsx(path)
        records = rows_to_records(rows)
        mapped = [apply_defaults(map_excel_record(record)) for record in records]
        print(json.dumps({'records': mapped, 'count': len(mapped)}, ensure_ascii=False))
        return

    if mode == 'contract':
        ext = path.suffix.lower()
        if ext == '.docx':
            data = parse_docx(path)
        elif ext == '.pdf':
            data = parse_pdf(path)
        else:
            data = parse_text(path)
        print(json.dumps({'record': apply_defaults(data)}, ensure_ascii=False))
        return

    if mode == 'ocr':
        ext = path.suffix.lower()
        if ext == '.pdf':
            data = parse_pdf(path)
        else:
            data = parse_image(path)
        print(json.dumps({'record': apply_defaults(data)}, ensure_ascii=False))
        return

    print(json.dumps({'error': f'unknown mode: {mode}'}, ensure_ascii=False))
    sys.exit(1)


if __name__ == '__main__':
    try:
        main()
    except Exception as exc:
        print(json.dumps({'error': str(exc)}, ensure_ascii=False))
        sys.exit(1)
